from docx import Document
from docx.shared import Inches,Cm

document = Document()

sections = document.sections
for section in sections:
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)

p = document.add_paragraph()
r = p.add_run()
r.add_picture('curse.jpg',width=Inches(2.5), height=Inches(3.5))
r.add_picture('curse2.jpg',width=Inches(2.5), height=Inches(3.5))

document.save('demo.docx')
